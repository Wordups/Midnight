"""
╔══════════════════════════════════════════════════════════════════════╗
║     M I D N I G H T  —  Policy Migration Engine                      ║
║    Takeoff Product                                                   ║
║    UI: Palo Alto Networks design language                            ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import os
import tempfile
from datetime import datetime, date

import streamlit as st
from groq import Groq
from docx import Document
from hps_policy_migration_builder import build_policy_document

# ── rembg is optional — logo upload works without it ──────────────────────────
try:
    from rembg import remove as rembg_remove
    REMBG_OK = True
except Exception:
    REMBG_OK = False


# ═════════════════════════════════════════════════════════════════════════════
# CONFIG
# ═════════════════════════════════════════════════════════════════════════════
LOCAL_GROQ_API_KEY = ""   # ← paste key here for local testing only

TEMPLATE_OPTIONS = [
    "Generic Policy Template",
    "Wipro HealthPlan Services",
]

PAGES = ["Overview", "Migrate a Policy", "Create a Policy"]

EXTRACTION_PROMPT = """
You are a policy migration specialist.

Read the attached legacy policy document and extract ALL content into the exact
Python dictionary structure below.

STRICT RULES:
- Do NOT summarize, rewrite, or remove content
- Preserve source wording as closely as possible
- Fix only minor spacing / punctuation / obvious grammar issues
- Map all content into the correct field
- For procedure items use exactly one type:
  "para"           = standalone paragraph
  "heading"        = bold underlined subsection title
  "bullet"         = first-level bullet
  "sub-bullet"     = second-level bullet
  "bold_intro"     = paragraph starting with bold label; keys: "bold" and "rest"
  "bold_intro_semi"= same as bold_intro but "rest" contains semicolons
  "empty"          = blank spacer line

Return ONLY a valid Python dictionary assignment.
No explanation. No markdown fences. No preamble.
Start your response with:
POLICY_DATA = {
End with the closing brace.

POLICY_DATA = {
    "policy_name": "",
    "policy_number": "",
    "version": "",
    "grc_id": "",
    "supersedes": "",
    "effective_date": "",
    "last_reviewed": "",
    "last_revised": "",
    "custodians": "",
    "owner_name": "",
    "owner_title": "",
    "approver_name": "",
    "approver_title": "",
    "date_signed": "",
    "date_approved": "",
    "applicable_to": {
        "hps_inc": False,
        "agency": True,
        "corporate": True,
        "govt_affairs": False,
        "legal_review": False
    },
    "policy_types": {
        "carrier_specific": False,
        "cross_carrier": False,
        "global": False,
        "on_off_hix": False
    },
    "line_of_business": {
        "all_lobs": True,
        "specific_lob": "",
        "specific_lob_checked": False
    },
    "purpose": "",
    "definitions": {},
    "policy_statement": "",
    "procedures": [],
    "related_policies": [],
    "citations": [],
    "revision_history": []
}

HERE IS THE LEGACY POLICY DOCUMENT:
"""


# ═════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ═════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Midnight",
    page_icon="🌑",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# ═════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ═════════════════════════════════════════════════════════════════════════════
_defaults = {
    "page":                  "Overview",
    "selected_template":     TEMPLATE_OPTIONS[0],
    "migration_policy_data": None,
    "created_policy_data":   None,
    "logo_path":             None,
    "logo_preview_name":     "",
    "migration_filename":    None,
    "migration_docx":        None,
    "created_filename":      None,
    "created_docx":          None,
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ═════════════════════════════════════════════════════════════════════════════
# GLOBAL CSS  —  Palo Alto Networks design language
# ═════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
/* ── Reset ──────────────────────────────────────────────────────────────── */
html, body, [class*="css"] { background: transparent !important; }

/* ── Root tokens ─────────────────────────────────────────────────────────── */
:root {
    --bg:           #020c1b;
    --bg2:          #051525;
    --surface:      #0a1628;
    --surface2:     #0f1e35;
    --border:       rgba(255,255,255,0.09);
    --border-strong:rgba(255,255,255,0.16);
    --text:         #e8edf5;
    --text-muted:   #8a9bb5;
    --text-dim:     #4f6180;
    --cyan:         #00c3e3;
    --cyan-dim:     rgba(0,195,227,0.12);
    --orange:       #fa582d;
    --orange-dim:   rgba(250,88,45,0.12);
    --green:        #00d68f;
    --green-dim:    rgba(0,214,143,0.10);
    --radius:       8px;
    --radius-lg:    12px;
}

/* ── App shell ───────────────────────────────────────────────────────────── */
.stApp {
    background:
        radial-gradient(ellipse 70% 30% at 50% 0%, rgba(0,195,227,0.06), transparent),
        linear-gradient(180deg, #020c1b 0%, #030f1e 100%);
    color: var(--text);
    font-family: 'Inter', 'Segoe UI', system-ui, sans-serif;
}

.block-container {
    max-width: 1320px;
    padding-top: 0.75rem;
    padding-bottom: 2.5rem;
    background: transparent !important;
}

header { visibility: hidden; }

div[data-testid="stVerticalBlock"] > div {
    background: transparent !important;
}

/* ── Top nav bar ─────────────────────────────────────────────────────────── */
.topbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    padding: 0.85rem 1.4rem;
    margin-bottom: 1.25rem;
}

.brand-lockup { display: flex; flex-direction: column; }
.brand-eyebrow {
    color: var(--text-dim);
    font-size: 0.68rem;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    margin-bottom: 0.1rem;
}
.brand-name {
    color: #ffffff;
    font-size: 1.45rem;
    font-weight: 700;
    letter-spacing: 0.06em;
}
.brand-sub {
    color: var(--text-muted);
    font-size: 0.82rem;
    margin-top: 0.08rem;
}

/* Streamlit radio restyled as nav pills */
div[data-testid="stRadio"] > div {
    gap: 0.3rem;
    flex-direction: row;
}
div[data-testid="stRadio"] label {
    background: transparent !important;
    border: 1px solid transparent !important;
    border-radius: 6px !important;
    padding: 0.5rem 1rem !important;
    transition: all .15s ease;
    cursor: pointer;
}
div[data-testid="stRadio"] label:hover {
    background: rgba(255,255,255,0.05) !important;
    border-color: var(--border) !important;
}
div[data-testid="stRadio"] label p {
    color: var(--text-muted) !important;
    font-size: 0.88rem !important;
    font-weight: 500 !important;
}
div[data-testid="stRadio"] label:has(input:checked) {
    background: var(--cyan-dim) !important;
    border-color: var(--cyan) !important;
}
div[data-testid="stRadio"] label:has(input:checked) p {
    color: var(--cyan) !important;
    font-weight: 600 !important;
}

/* ── Hero ────────────────────────────────────────────────────────────────── */
.hero {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    padding: 3rem 2.5rem;
    margin-bottom: 1.25rem;
    position: relative;
    overflow: hidden;
}
.hero::after {
    content: "";
    position: absolute;
    top: -60px; right: -80px;
    width: 320px; height: 320px;
    background: radial-gradient(circle, rgba(0,195,227,0.09), transparent 65%);
    pointer-events: none;
}

.hero-eyebrow {
    color: var(--cyan);
    font-size: 0.72rem;
    letter-spacing: 0.24em;
    text-transform: uppercase;
    font-weight: 600;
    margin-bottom: 1rem;
}
.hero-headline {
    color: #ffffff;
    font-size: 3rem;
    font-weight: 700;
    line-height: 1.1;
    margin-bottom: 1rem;
    letter-spacing: -0.01em;
}
.hero-headline .hl { color: var(--cyan); }
.hero-body {
    color: var(--text-muted);
    font-size: 1.02rem;
    line-height: 1.7;
    max-width: 640px;
    margin-bottom: 1.5rem;
}

/* ── Cards / Panels ──────────────────────────────────────────────────────── */
.card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    padding: 1.25rem;
}

.card-title {
    color: #ffffff;
    font-size: 1.1rem;
    font-weight: 700;
    margin-bottom: 0.35rem;
}
.card-body {
    color: var(--text-muted);
    font-size: 0.9rem;
    line-height: 1.65;
}

/* ── Stat chips ──────────────────────────────────────────────────────────── */
.stat-row { display: flex; gap: 0.8rem; flex-wrap: wrap; margin-top: 1.2rem; }
.stat {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 0.85rem 1rem;
    min-width: 130px;
}
.stat-num {
    color: var(--cyan);
    font-size: 1.85rem;
    font-weight: 700;
    line-height: 1;
    margin-bottom: 0.25rem;
}
.stat-lbl { color: var(--text-muted); font-size: 0.8rem; }

/* ── Feature grid ────────────────────────────────────────────────────────── */
.feature-grid {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 1rem;
    margin-bottom: 1.25rem;
}
.feature {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    padding: 1.1rem;
}
.feature-icon {
    color: var(--cyan);
    font-size: 1.15rem;
    margin-bottom: 0.6rem;
}
.feature-title {
    color: #ffffff;
    font-size: 0.98rem;
    font-weight: 600;
    margin-bottom: 0.3rem;
}
.feature-copy {
    color: var(--text-muted);
    font-size: 0.86rem;
    line-height: 1.6;
}

/* ── Workspace ───────────────────────────────────────────────────────────── */
.ws-title {
    color: #ffffff;
    font-size: 1.6rem;
    font-weight: 700;
    margin-bottom: 0.2rem;
}
.ws-sub {
    color: var(--text-muted);
    font-size: 0.9rem;
    margin-bottom: 1rem;
}

/* ── Info / note boxes ───────────────────────────────────────────────────── */
.note {
    background: var(--surface2);
    border-left: 3px solid var(--border-strong);
    border-radius: 0 var(--radius) var(--radius) 0;
    padding: 0.75rem 1rem;
    color: var(--text-muted);
    font-size: 0.86rem;
    line-height: 1.6;
    margin: 0.6rem 0;
}
.note strong { color: var(--text); }

.success-banner {
    background: var(--green-dim);
    border: 1px solid rgba(0,214,143,0.3);
    border-radius: var(--radius);
    padding: 0.75rem 1rem;
    color: #4df0b8;
    font-weight: 600;
    text-align: center;
    margin: 0.75rem 0;
}

.preview-pane {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    padding: 1.1rem;
}

.caption {
    color: var(--text-dim);
    font-size: 0.78rem;
    margin-bottom: 0.5rem;
}

/* ── Inputs ──────────────────────────────────────────────────────────────── */
.stButton > button {
    width: 100% !important;
    background: var(--cyan) !important;
    color: #000000 !important;
    border: none !important;
    border-radius: var(--radius) !important;
    padding: 0.78rem 1rem !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    letter-spacing: 0.03em !important;
    transition: opacity .15s ease, transform .12s ease !important;
}
.stButton > button:hover {
    opacity: 0.88 !important;
    transform: translateY(-1px) !important;
}
.stButton > button p,
.stButton > button span { color: #000000 !important; }

.stDownloadButton > button {
    width: 100% !important;
    background: var(--surface2) !important;
    color: var(--text) !important;
    border: 1px solid var(--border-strong) !important;
    border-radius: var(--radius) !important;
    padding: 0.78rem 1rem !important;
    font-weight: 600 !important;
}
.stDownloadButton > button:hover {
    border-color: var(--cyan) !important;
    color: var(--cyan) !important;
}

.stTextInput input,
.stTextArea textarea {
    background: var(--surface2) !important;
    color: var(--text) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius) !important;
}
.stTextInput input:focus,
.stTextArea textarea:focus {
    border-color: var(--cyan) !important;
    box-shadow: 0 0 0 2px rgba(0,195,227,0.15) !important;
}

.stSelectbox div[data-baseweb="select"] > div {
    background: var(--surface2) !important;
    color: var(--text) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius) !important;
}

[data-testid="stWidgetLabel"] p {
    color: var(--text) !important;
    font-weight: 500 !important;
    font-size: 0.88rem !important;
}

input::placeholder, textarea::placeholder {
    color: var(--text-dim) !important;
    opacity: 1 !important;
}

.stFileUploader > div {
    background: var(--surface2) !important;
    border: 1px dashed var(--border-strong) !important;
    border-radius: var(--radius) !important;
}

.stProgress > div > div {
    background: var(--cyan) !important;
}

.stAlert {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    color: var(--text) !important;
    border-radius: var(--radius) !important;
}

h1, h2, h3, h4 { color: #ffffff !important; }
p, span        { color: var(--text-muted); }

@media (max-width: 900px) {
    .feature-grid { grid-template-columns: 1fr; }
    .hero-headline { font-size: 2.1rem; }
    .hero { padding: 2rem 1.25rem; }
}
</style>
""", unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═════════════════════════════════════════════════════════════════════════════

def get_api_key() -> str:
    try:
        secret = st.secrets.get("GROQ_API_KEY", "")
    except Exception:
        secret = ""
    return secret or os.getenv("GROQ_API_KEY", "") or LOCAL_GROQ_API_KEY


def parse_policy_data(raw: str):
    if "POLICY_DATA = {" in raw:
        raw = raw[raw.index("POLICY_DATA = {"):]
    raw = raw.replace("\u201c", '"').replace("\u201d", '"').replace("\u2019", "'")
    ns = {}
    exec(raw, {}, ns)
    return ns.get("POLICY_DATA")


def extract_text_from_docx(uploaded_file) -> str:
    doc = Document(uploaded_file)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                ct = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if ct:
                    row_parts.append(ct)
            if row_parts:
                lines.append(" | ".join(row_parts))
    return "\n".join(lines)


def get_uploaded_text(uploaded_file) -> str:
    if uploaded_file.name.lower().endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    return uploaded_file.read().decode("utf-8", errors="ignore")


def norm_date(value: str) -> str:
    value = str(value).strip()
    if not value:
        return ""
    for fmt in ("%m/%d/%Y", "%m/%d/%y", "%Y-%m-%d"):
        try:
            return datetime.strptime(value, fmt).strftime("%-m/%-d/%Y")
        except Exception:
            pass
    return value


def parse_date_safe(value: str):
    v = norm_date(value)
    if not v:
        return None
    try:
        return datetime.strptime(v, "%m/%d/%Y")
    except Exception:
        return None


def fallback(current: str, source: str) -> str:
    return current if str(current).strip() else source


def validate_dates(eff, revw, revd, signed, approved):
    errs = []
    e = parse_date_safe(eff)
    rw = parse_date_safe(revw)
    rd = parse_date_safe(revd)
    s  = parse_date_safe(signed)
    a  = parse_date_safe(approved)
    if e and rw and rw < e:    errs.append("Last Reviewed cannot be earlier than Effective Date.")
    if e and rd and rd < e:    errs.append("Last Revised cannot be earlier than Effective Date.")
    if s and a  and a  < s:    errs.append("Date Approved cannot be earlier than Date Signed.")
    return errs


def split_lines(text: str):
    return [l.strip() for l in str(text).splitlines() if l.strip()]


def procedures_from_text(text: str):
    procs = []
    for line in split_lines(text):
        if line.startswith("- "):
            procs.append({"type": "bullet", "text": line[2:].strip()})
        else:
            procs.append({"type": "para", "text": line})
    return procs


def build_creation_data(fields: dict, template_name: str) -> dict:
    eff    = norm_date(fields["effective_date"])
    revw   = norm_date(fallback(fields["last_reviewed"],  eff))
    revd   = norm_date(fallback(fields["last_revised"],   eff))
    signed = norm_date(fallback(fields["date_signed"],    eff))
    approv = norm_date(fallback(fields["date_approved"],  signed))

    defs = {}
    for line in split_lines(fields["definitions_text"]):
        if ":" in line:
            k, v = line.split(":", 1)
            defs[k.strip()] = v.strip()
        else:
            defs[line.strip()] = ""

    return {
        "policy_name":    fields["policy_name"],
        "policy_number":  fields["policy_number"],
        "version":        fields["version"] or "V1.0",
        "grc_id":         fields["grc_id"],
        "supersedes":     fields["supersedes"],
        "effective_date": eff,
        "last_reviewed":  revw,
        "last_revised":   revd,
        "custodians":     fields["custodians"],
        "owner_name":     fields["owner_name"],
        "owner_title":    fields["owner_title"],
        "approver_name":  fields["approver_name"],
        "approver_title": fields["approver_title"],
        "date_signed":    signed,
        "date_approved":  approv,
        "applicable_to":  {"hps_inc": False, "agency": True, "corporate": True,
                           "govt_affairs": False, "legal_review": False},
        "policy_types":   {"carrier_specific": False, "cross_carrier": False,
                           "global": template_name == "Generic Policy Template",
                           "on_off_hix": False},
        "line_of_business": {"all_lobs": True, "specific_lob": "", "specific_lob_checked": False},
        "purpose":         fields["purpose"],
        "definitions":     defs,
        "policy_statement":fields["policy_statement"],
        "procedures":      procedures_from_text(fields["procedures_text"]),
        "related_policies":split_lines(fields["related_policies_text"]),
        "citations":       split_lines(fields["citations_text"]),
        "revision_history":[],
        "template_name":   template_name,
    }


def save_logo(uploaded_file) -> str:
    """Save logo; remove background if rembg available, else save as-is."""
    os.makedirs("assets", exist_ok=True)
    os.makedirs("uploads", exist_ok=True)

    stem = "".join(c for c in os.path.splitext(uploaded_file.name)[0]
                   if c.isalnum() or c in "-_") or "logo"
    ext  = os.path.splitext(uploaded_file.name)[1].lower()
    raw_path = f"uploads/{stem}{ext}"
    out_path = f"assets/{stem}_processed.png"

    raw_bytes = uploaded_file.getbuffer()
    with open(raw_path, "wb") as f:
        f.write(raw_bytes)

    if REMBG_OK and ext != ".png":
        with open(raw_path, "rb") as f:
            result = rembg_remove(f.read())
        with open(out_path, "wb") as f:
            f.write(result)
    else:
        # PNG doesn't need background removal; non-rembg env: use as-is
        import shutil
        shutil.copy(raw_path, out_path)

    return out_path


def render_logo_ui(key_suffix: str):
    st.markdown('<div class="caption">Optional — upload a logo for the document header.</div>',
                unsafe_allow_html=True)
    logo_file = st.file_uploader(
        "Logo",
        type=["png", "jpg", "jpeg", "webp"],
        key=f"logo_{key_suffix}",
        label_visibility="collapsed",
    )
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Process Logo", key=f"process_{key_suffix}"):
            if not logo_file:
                st.error("Upload a logo first.")
            else:
                try:
                    path = save_logo(logo_file)
                    st.session_state["logo_path"] = path
                    st.session_state["logo_preview_name"] = logo_file.name
                    st.success("Logo ready.")
                except Exception as e:
                    st.error(f"Logo error: {e}")
    with c2:
        if st.button("Clear Logo", key=f"clear_{key_suffix}"):
            st.session_state["logo_path"] = None
            st.session_state["logo_preview_name"] = ""
            st.success("Cleared.")

    lp = st.session_state.get("logo_path")
    if lp and os.path.exists(lp):
        st.image(lp, caption=st.session_state.get("logo_preview_name", "Logo"),
                 use_container_width=True)


def render_preview(policy_data: dict):
    st.markdown('<div class="preview-pane">', unsafe_allow_html=True)
    st.markdown(f"### {policy_data.get('policy_name', '—')}")
    st.markdown(f"**{policy_data.get('policy_number','')}  ·  "
                f"v{policy_data.get('version','')}  ·  "
                f"{policy_data.get('effective_date','')}**")
    st.divider()
    st.markdown(f"**Owner:** {policy_data.get('owner_name','')} — {policy_data.get('owner_title','')}")
    st.markdown(f"**Approver:** {policy_data.get('approver_name','')} — {policy_data.get('approver_title','')}")

    purpose = policy_data.get("purpose","")
    if purpose:
        st.markdown("#### Purpose")
        st.write(purpose[:600] + ("…" if len(purpose) > 600 else ""))

    stmt = policy_data.get("policy_statement","")
    if stmt:
        st.markdown("#### Policy Statement")
        st.write(stmt[:400] + ("…" if len(stmt) > 400 else ""))

    defs = policy_data.get("definitions") or {}
    if defs:
        st.markdown("#### Definitions")
        for k, v in list(defs.items())[:5]:
            st.markdown(f"- **{k}:** {v}")
        if len(defs) > 5:
            st.caption(f"+ {len(defs)-5} more definitions")

    procs = policy_data.get("procedures") or []
    if procs:
        st.markdown("#### Procedures")
        for item in procs[:8]:
            kind = item.get("type","")
            text = item.get("text","") or item.get("rest","")
            if kind == "bullet":      st.markdown(f"- {text}")
            elif kind == "sub-bullet":st.markdown(f"  - {text}")
            elif kind == "heading":   st.markdown(f"**{text}**")
            elif kind in ("bold_intro","bold_intro_semi"):
                st.markdown(f"**{item.get('bold','')}** {item.get('rest','')}")
            elif text:                st.write(text)
        if len(procs) > 8:
            st.caption(f"+ {len(procs)-8} more items")

    rev = policy_data.get("revision_history") or []
    if rev:
        st.markdown(f"#### Revision History  ({len(rev)} entries)")

    st.markdown("</div>", unsafe_allow_html=True)


def build_doc(policy_data: dict) -> tuple[str, bytes]:
    name   = policy_data.get("policy_name",   "Policy")
    number = policy_data.get("policy_number", "SEC-P")
    ver    = policy_data.get("version",       "V1.0")
    fname  = f"{number} {name} {ver}-NEW.docx"

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = tmp.name

    build_policy_document(
        policy_data,
        tmp_path,
        logo_path=st.session_state.get("logo_path"),
    )

    with open(tmp_path, "rb") as f:
        docx_bytes = f.read()

    return fname, docx_bytes


def run_extraction(source_text: str, template_name: str) -> dict | None:
    api_key = get_api_key()
    if not api_key:
        st.error("No Groq API key. Set GROQ_API_KEY in Streamlit secrets or line 25 for local testing.")
        return None

    prog   = st.progress(0)
    status = st.empty()

    try:
        status.markdown('<div class="caption">Reading document…</div>', unsafe_allow_html=True)
        prog.progress(15)

        client = Groq(api_key=api_key)

        status.markdown('<div class="caption">Extracting policy structure…</div>', unsafe_allow_html=True)
        prog.progress(35)

        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": EXTRACTION_PROMPT + "\n\n" + source_text}],
            temperature=0.1,
            max_tokens=8000,
        )

        raw = resp.choices[0].message.content.strip()

        status.markdown('<div class="caption">Parsing extracted data…</div>', unsafe_allow_html=True)
        prog.progress(65)

        policy_data = parse_policy_data(raw)

        if not policy_data:
            st.error("Model response could not be parsed into POLICY_DATA.")
            with st.expander("Raw model output"):
                st.code(raw)
            return None

        policy_data["template_name"] = template_name
        prog.progress(100)
        status.empty()
        return policy_data

    except Exception as e:
        prog.empty()
        status.empty()
        st.error(f"Extraction failed: {e}")
        return None


# ═════════════════════════════════════════════════════════════════════════════
# TOP NAV
# ═════════════════════════════════════════════════════════════════════════════
nav_left, nav_right = st.columns([0.6, 0.4])

with nav_left:
    st.markdown("""
    <div class="topbar">
        <div class="brand-lockup">
            <div class="brand-eyebrow">Takeoff Product</div>
            <div class="brand-name">MIDNIGHT</div>
            <div class="brand-sub">Policy Migration Engine</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

with nav_right:
    st.markdown('<div style="height:0.75rem"></div>', unsafe_allow_html=True)
    selected = st.radio(
        "nav",
        PAGES,
        horizontal=True,
        label_visibility="collapsed",
        index=PAGES.index(st.session_state["page"]),
        key="top_nav",
    )
    st.session_state["page"] = selected


# ═════════════════════════════════════════════════════════════════════════════
# PAGE: OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
if st.session_state["page"] == "Overview":

    # Hero
    st.markdown("""
    <div class="hero">
        <div class="hero-eyebrow">Policy intelligence engine</div>
        <div class="hero-headline">
            Migrate. Create.<br><span class="hl">Ship cleaner policy.</span>
        </div>
        <div class="hero-body">
            Midnight converts legacy documents and structured intake into
            audit-ready, template-faithful Word output — without the manual effort.
        </div>
    </div>
    """, unsafe_allow_html=True)

    cta1, cta2, _, _ = st.columns([0.18, 0.18, 0.32, 0.32])
    with cta1:
        if st.button("Migrate a Policy", key="hero_migrate"):
            st.session_state["page"] = "Migrate a Policy"
            st.rerun()
    with cta2:
        if st.button("Create a Policy", key="hero_create"):
            st.session_state["page"] = "Create a Policy"
            st.rerun()

    st.markdown('<div style="height:0.75rem"></div>', unsafe_allow_html=True)

    # Feature grid
    st.markdown("""
    <div class="feature-grid">
        <div class="feature">
            <div class="feature-icon">⬆</div>
            <div class="feature-title">Migrate Policy</div>
            <div class="feature-copy">Upload a legacy document (.docx, .txt, .md) and
            convert it into the selected template using AI extraction.</div>
        </div>
        <div class="feature">
            <div class="feature-icon">✦</div>
            <div class="feature-title">Create Policy</div>
            <div class="feature-copy">Build a new policy from structured intake — smart
            date defaults, inline preview, controlled final output.</div>
        </div>
        <div class="feature">
            <div class="feature-icon">⬇</div>
            <div class="feature-title">Download Ready</div>
            <div class="feature-copy">Review the extracted data, then generate a
            pixel-faithful HPS-template Word document in one click.</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Bottom row
    bl, br = st.columns([0.55, 0.45], gap="large")

    with bl:
        st.markdown("""
        <div class="card" style="background:var(--surface2);">
            <div class="card-title">Built for policy operations teams</div>
            <div class="card-body">
                Manual policy updates create inconsistency and slow audit prep.
                Midnight provides a controlled path from source content to final document —
                extraction, preview, and generation in one workflow.
            </div>
            <div class="stat-row">
                <div class="stat">
                    <div class="stat-num">80h</div>
                    <div class="stat-lbl">Reducible manual effort across a typical backlog</div>
                </div>
                <div class="stat">
                    <div class="stat-num">1</div>
                    <div class="stat-lbl">Structured workflow from source to final output</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    with br:
        st.markdown("""
        <div class="card">
            <div class="card-title">How it works</div>
            <div class="card-body">
                Select a template, choose a workflow, upload or enter content,
                review the live preview, then generate the final document.<br><br>
                Supports .docx, .txt, and .md source files. Revision history,
                definitions, and all metadata carry through automatically.
            </div>
        </div>
        """, unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
# PAGE: MIGRATE A POLICY
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state["page"] == "Migrate a Policy":

    st.markdown('<div class="ws-title">Migrate a Policy</div>', unsafe_allow_html=True)
    st.markdown('<div class="ws-sub">Convert an existing document into a structured template.</div>',
                unsafe_allow_html=True)

    settings_col, upload_col = st.columns([0.32, 0.68], gap="medium")

    with settings_col:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">Settings</div>', unsafe_allow_html=True)

        sel_tmpl = st.selectbox(
            "Target Template",
            TEMPLATE_OPTIONS,
            index=TEMPLATE_OPTIONS.index(st.session_state["selected_template"]),
            key="migrate_template",
        )
        st.session_state["selected_template"] = sel_tmpl

        st.markdown("""
        <div class="note">
            <strong>Workflow:</strong> Upload → Transform → Review preview → Generate final document.
        </div>
        """, unsafe_allow_html=True)

        render_logo_ui("migrate")
        st.markdown('</div>', unsafe_allow_html=True)

    with upload_col:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">Upload Legacy Policy</div>', unsafe_allow_html=True)
        st.markdown('<div class="caption">Supported: .docx, .txt, .md — up to 200 MB</div>',
                    unsafe_allow_html=True)

        uploaded = st.file_uploader(
            "upload",
            type=["docx", "txt", "md"],
            label_visibility="collapsed",
            key="migrate_upload",
        )

        st.markdown("""
        <div class="note">
            <strong>System note:</strong> API access is configured on the backend via
            Streamlit secrets or environment variables.
        </div>
        """, unsafe_allow_html=True)

        run_btn = st.button("Transform Policy", key="run_migrate")
        st.markdown('</div>', unsafe_allow_html=True)

    # Preview panel
    st.markdown('<div style="height:0.75rem"></div>', unsafe_allow_html=True)
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">Preview</div>', unsafe_allow_html=True)

    if st.session_state["migration_policy_data"]:
        render_preview(st.session_state["migration_policy_data"])
    else:
        st.info("Upload a document and run Transform to preview the extracted policy.")

    st.markdown('</div>', unsafe_allow_html=True)

    # Run extraction
    if run_btn:
        if not uploaded:
            st.error("Please upload a legacy policy document.")
            st.stop()
        src_text = get_uploaded_text(uploaded)
        if len(src_text.strip()) < 50:
            st.error("Document appears empty or too short.")
            st.stop()
        pd = run_extraction(src_text, st.session_state["selected_template"])
        if pd:
            st.session_state["migration_policy_data"] = pd
            # Clear previous output when re-running
            st.session_state["migration_docx"] = None
            st.session_state["migration_filename"] = None
            st.rerun()

    # Generate + download
    if st.session_state["migration_policy_data"]:
        st.markdown('<div style="height:0.5rem"></div>', unsafe_allow_html=True)
        _, gc, _ = st.columns([1, 1.4, 1])
        with gc:
            if st.button("Generate Final Document", key="gen_migrate"):
                try:
                    fname, docx_bytes = build_doc(st.session_state["migration_policy_data"])
                    st.session_state["migration_filename"] = fname
                    st.session_state["migration_docx"]     = docx_bytes
                except Exception as e:
                    st.error(f"Build failed: {e}")

            if st.session_state.get("migration_docx"):
                st.markdown('<div class="success-banner">✓ Transformation complete</div>',
                            unsafe_allow_html=True)
                st.download_button(
                    label=f"↓ Download {st.session_state['migration_filename']}",
                    data=st.session_state["migration_docx"],
                    file_name=st.session_state["migration_filename"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_migrate",
                )


# ═════════════════════════════════════════════════════════════════════════════
# PAGE: CREATE A POLICY
# ═════════════════════════════════════════════════════════════════════════════
else:

    st.markdown('<div class="ws-title">Create a Policy</div>', unsafe_allow_html=True)
    st.markdown('<div class="ws-sub">Generate a new policy from structured intake.</div>',
                unsafe_allow_html=True)

    settings_col, form_col = st.columns([0.32, 0.68], gap="medium")

    with settings_col:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">Settings</div>', unsafe_allow_html=True)

        sel_tmpl = st.selectbox(
            "Target Template",
            TEMPLATE_OPTIONS,
            index=TEMPLATE_OPTIONS.index(st.session_state["selected_template"]),
            key="create_template",
        )
        st.session_state["selected_template"] = sel_tmpl

        st.markdown("""
        <div class="note">
            Blank date fields auto-fill from the Effective Date.
            All fields can be edited after preview.
        </div>
        """, unsafe_allow_html=True)

        render_logo_ui("create")
        st.markdown('</div>', unsafe_allow_html=True)

    with form_col:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">Intake</div>', unsafe_allow_html=True)

        with st.form("create_form"):
            c1, c2 = st.columns(2)
            with c1:
                policy_name   = st.text_input("Policy Name",   value="")
                policy_number = st.text_input("Policy Number", value="")
                version       = st.text_input("Version",       value="V1.0")
                grc_id        = st.text_input("GRC ID",        value="")
                effective_date= st.text_input("Effective Date",
                                              value=date.today().strftime("%-m/%-d/%Y"))
                last_reviewed = st.text_input("Last Reviewed", value="",
                                              placeholder="Defaults to Effective Date")
                last_revised  = st.text_input("Last Revised",  value="",
                                              placeholder="Defaults to Effective Date")
                supersedes    = st.text_input("Supersedes",    value="")
            with c2:
                custodians    = st.text_input("Custodians",    value="")
                owner_name    = st.text_input("Owner Name",    value="")
                owner_title   = st.text_input("Owner Title",   value="")
                approver_name = st.text_input("Approver Name", value="")
                approver_title= st.text_input("Approver Title",value="")
                date_signed   = st.text_input("Date Signed",   value="",
                                              placeholder="Defaults to Effective Date")
                date_approved = st.text_input("Date Approved", value="",
                                              placeholder="Defaults to Date Signed")

            purpose            = st.text_area("Purpose and Scope",        height=110)
            definitions_text   = st.text_area("Definitions  (Term: Definition, one per line)", height=100)
            policy_statement   = st.text_area("Policy Statement",          height=120)
            procedures_text    = st.text_area("Procedures  (use '- ' for bullets)", height=180)
            related_policies_text = st.text_area("Related Policies  (one per line)", height=90)
            citations_text     = st.text_area("Citations / References  (one per line)", height=90)

            build_preview_btn = st.form_submit_button("Build Preview")

        st.markdown('</div>', unsafe_allow_html=True)

    # Smart defaults display + preview
    preview_dates = {
        "last_reviewed": norm_date(fallback(last_reviewed, effective_date)),
        "last_revised":  norm_date(fallback(last_revised,  effective_date)),
        "date_signed":   norm_date(fallback(date_signed,   effective_date)),
        "date_approved": norm_date(fallback(date_approved, fallback(date_signed, effective_date))),
    }

    st.markdown('<div style="height:0.75rem"></div>', unsafe_allow_html=True)
    dl_col, pv_col = st.columns([0.4, 0.6], gap="medium")

    with dl_col:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">Smart Defaults</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="note">
            <strong>Last Reviewed:</strong> {preview_dates["last_reviewed"] or "—"}<br>
            <strong>Last Revised:</strong>  {preview_dates["last_revised"]  or "—"}<br>
            <strong>Date Signed:</strong>   {preview_dates["date_signed"]   or "—"}<br>
            <strong>Date Approved:</strong> {preview_dates["date_approved"] or "—"}
        </div>
        """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with pv_col:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">Preview</div>', unsafe_allow_html=True)
        if st.session_state["created_policy_data"]:
            render_preview(st.session_state["created_policy_data"])
        else:
            st.info("Complete the intake form and click Build Preview.")
        st.markdown('</div>', unsafe_allow_html=True)

    # Handle form submission
    if build_preview_btn:
        errs = validate_dates(
            effective_date,
            preview_dates["last_reviewed"],
            preview_dates["last_revised"],
            preview_dates["date_signed"],
            preview_dates["date_approved"],
        )
        if errs:
            for e in errs:
                st.error(e)
        else:
            pd = build_creation_data({
                "policy_name": policy_name, "policy_number": policy_number,
                "version": version, "grc_id": grc_id, "supersedes": supersedes,
                "effective_date": effective_date, "last_reviewed": last_reviewed,
                "last_revised": last_revised, "custodians": custodians,
                "owner_name": owner_name, "owner_title": owner_title,
                "approver_name": approver_name, "approver_title": approver_title,
                "date_signed": date_signed, "date_approved": date_approved,
                "purpose": purpose, "definitions_text": definitions_text,
                "policy_statement": policy_statement, "procedures_text": procedures_text,
                "related_policies_text": related_policies_text, "citations_text": citations_text,
            }, st.session_state["selected_template"])
            st.session_state["created_policy_data"] = pd
            st.session_state["created_docx"] = None
            st.session_state["created_filename"] = None
            st.rerun()

    # Generate + download
    if st.session_state["created_policy_data"]:
        st.markdown('<div style="height:0.5rem"></div>', unsafe_allow_html=True)
        _, gc, _ = st.columns([1, 1.4, 1])
        with gc:
            if st.button("Generate Policy Document", key="gen_create"):
                try:
                    fname, docx_bytes = build_doc(st.session_state["created_policy_data"])
                    st.session_state["created_filename"] = fname
                    st.session_state["created_docx"]     = docx_bytes
                except Exception as e:
                    st.error(f"Build failed: {e}")

            if st.session_state.get("created_docx"):
                st.markdown('<div class="success-banner">✓ Policy generated</div>',
                            unsafe_allow_html=True)
                st.download_button(
                    label=f"↓ Download {st.session_state['created_filename']}",
                    data=st.session_state["created_docx"],
                    file_name=st.session_state["created_filename"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_create",
                )
