"""
================================================================================
HPS SEC-P Policy Template Generator  —  Blank Template
================================================================================
Generates a blank SEC-P policy document matching the Wipro HealthPlan Services
template exactly as shown in the official screenshots.

Run:
    python hps_blank_template.py

Output:
    SEC-P_Template_Master-BLANK.docx
================================================================================
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ─────────────────────────────────────────────────────────────────────
GRAY_LOGO    = "BFBFBF"   # Wipro logo banner background
GRAY_LABEL   = "D9D9D9"   # Label cells (Policy Name, Policy Number, etc.)
GRAY_SUBHDR  = "BFBFBF"   # Policy Owner/Approver bar & Policy Types/LOB headers
GRAY_SECTION = "D9D9D9"   # Section heading rows
WHITE        = "FFFFFF"
BLACK        = "000000"
WIPRO_RED    = "C00000"   # "wipro:" brand red
WIPRO_TEAL   = "17375E"   # "healthplan services" dark teal

# ── Dimensions ─────────────────────────────────────────────────────────────────
PAGE_W = int((8.5 - 0.75 - 0.75) * 1440)   # usable width in twips (0.75" margins)

META_L   = int(PAGE_W * 0.18)               # metadata left label col
META_MID = int(PAGE_W * 0.32)               # metadata left value col
META_RL  = int(PAGE_W * 0.18)               # metadata right label col
META_RV  = PAGE_W - META_L - META_MID - META_RL  # metadata right value col

APP_L    = int(PAGE_W * 0.23)               # "Applicable To" label col
APP_R    = PAGE_W - APP_L                   # checkbox item col

REV_C1   = int(PAGE_W * 0.12)              # Date
REV_C2   = int(PAGE_W * 0.13)              # Version Number
REV_C3   = int(PAGE_W * 0.20)              # Updated By
REV_C4   = PAGE_W - REV_C1 - REV_C2 - REV_C3  # Description of Update


# ══════════════════════════════════════════════════════════════════════════════
# XML / STYLE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def set_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill.upper())
    tcPr.append(shd)


def set_borders(cell, color="000000", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def set_col_width(cell, twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_row_height(row, twips, rule="atLeast"):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(twips)))
    trH.set(qn("w:hRule"), rule)
    trPr.append(trH)


def set_spacing(para, before=40, after=40):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    pPr.append(spc)


def rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def run(para, text, bold=False, italic=False, size=9.5,
        color=BLACK, font="Arial", underline=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    return r


def style_cell(cell, shade=WHITE, width=None):
    """Apply shading, borders, and optional column width to a cell."""
    set_shading(cell, shade)
    set_borders(cell)
    set_margins(cell)
    if width:
        set_col_width(cell, width)


def empty_cell(cell, shade=WHITE, width=None, height_rows=1):
    """Blank content cell with optional minimum height."""
    style_cell(cell, shade, width)
    cell.text = ""
    set_spacing(cell.paragraphs[0], 0, 0)


def label_cell(cell, text, width=None):
    """Gray label cell — right-aligned bold text."""
    style_cell(cell, GRAY_LABEL, width)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(p, 50, 50)
    run(p, text, bold=True, size=9.0)


def value_cell(cell, width=None):
    """White value cell — empty, ready for content."""
    style_cell(cell, WHITE, width)
    cell.text = ""
    set_spacing(cell.paragraphs[0], 50, 50)


def subheader_cell(cell, text, width=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Medium-gray sub-header cell (Policy Owner/Approver, Policy Types, LOB)."""
    style_cell(cell, GRAY_SUBHDR, width)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_spacing(p, 55, 55)
    run(p, text, bold=True, size=9.5)


def section_head_cell(cell, text, width=None):
    """Light-gray section heading cell — bold left-aligned."""
    style_cell(cell, GRAY_SECTION, width)
    set_margins(cell, top=60, bottom=60, left=80, right=80)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 60, 60)
    run(p, text, bold=True, size=10.0)


def set_table_width(tbl, total_twips, col_widths):
    """Set total table width and individual column widths."""
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    tblGrid = tbl._tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl._tbl.insert(0, tblGrid)
    for old in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(old)
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)


def gap(doc, size=80):
    """Small spacer paragraph between tables."""
    p = doc.add_paragraph()
    set_spacing(p, size, 0)


# ══════════════════════════════════════════════════════════════════════════════
# TEMPLATE BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_blank_template(output_path="SEC-P_Template_Master-BLANK.docx"):

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.9)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 1 — Wipro Logo Banner
    # ──────────────────────────────────────────────────────────────────────
    t_logo = doc.add_table(rows=1, cols=1)
    t_logo.style = "Table Grid"
    set_table_width(t_logo, PAGE_W, [PAGE_W])

    c = t_logo.cell(0, 0)
    set_shading(c, GRAY_LOGO)
    set_borders(c)
    set_margins(c, top=180, bottom=180, left=120, right=120)
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 0, 0)
    run(p, "wipro",                bold=True, size=22, color=WIPRO_RED)
    run(p, ":",                    bold=True, size=22, color=WIPRO_RED)
    run(p, "  healthplan services",           size=20, color=WIPRO_TEAL)

    gap(doc)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 2 — Metadata (Policy Name → Date Approved)
    # ──────────────────────────────────────────────────────────────────────
    t_meta = doc.add_table(rows=11, cols=4)
    t_meta.style = "Table Grid"
    set_table_width(t_meta, PAGE_W, [META_L, META_MID, META_RL, META_RV])

    # Row 0 — Policy Name (label | value spanning 3 cols)
    r = t_meta.rows[0]
    label_cell(r.cells[0], "Policy Name",   META_L)
    r.cells[1].merge(r.cells[3])
    value_cell(r.cells[1], PAGE_W - META_L)

    # Row 1 — Policy Number | Version Number
    r = t_meta.rows[1]
    label_cell(r.cells[0], "Policy Number",  META_L)
    value_cell(r.cells[1],                   META_MID)
    label_cell(r.cells[2], "Version Number", META_RL)
    value_cell(r.cells[3],                   META_RV)

    # Row 2 — (blank) | GRC ID Number
    r = t_meta.rows[2]
    style_cell(r.cells[0], GRAY_LABEL, META_L)
    r.cells[0].text = ""
    set_spacing(r.cells[0].paragraphs[0], 50, 50)
    value_cell(r.cells[1], META_MID)
    label_cell(r.cells[2], "GRC ID Number", META_RL)
    value_cell(r.cells[3],                  META_RV)

    # Row 3 — Supersedes Policy | Effective Date
    r = t_meta.rows[3]
    label_cell(r.cells[0], "Supersedes Policy", META_L)
    value_cell(r.cells[1],                       META_MID)
    label_cell(r.cells[2], "Effective Date",     META_RL)
    value_cell(r.cells[3],                       META_RV)

    # Row 4 — Last Reviewed Date | Last Revised Date
    r = t_meta.rows[4]
    label_cell(r.cells[0], "Last Reviewed Date", META_L)
    value_cell(r.cells[1],                        META_MID)
    label_cell(r.cells[2], "Last Revised Date",   META_RL)
    value_cell(r.cells[3],                        META_RV)

    # Row 5 — Policy Custodian Name(s)  (pre-filled per template standard)
    r = t_meta.rows[5]
    label_cell(r.cells[0], "Policy Custodian\nName(s)", META_L)
    r.cells[1].merge(r.cells[3])
    style_cell(r.cells[1], WHITE, PAGE_W - META_L)
    r.cells[1].text = ""
    p = r.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 50, 50)
    run(p, "Chelsea Sanchez, Alexis Taylor", size=9.5)

    # Row 6 — Policy Owner | Policy Approver sub-header bar
    r = t_meta.rows[6]
    r.cells[0].merge(r.cells[1])
    r.cells[2].merge(r.cells[3])
    subheader_cell(r.cells[0], "Policy Owner",    META_L + META_MID)
    subheader_cell(r.cells[2], "Policy Approver", META_RL + META_RV)

    # Row 7 — Name
    r = t_meta.rows[7]
    label_cell(r.cells[0], "Name", META_L)
    value_cell(r.cells[1],         META_MID)
    label_cell(r.cells[2], "Name", META_RL)
    value_cell(r.cells[3],         META_RV)

    # Row 8 — Title
    r = t_meta.rows[8]
    label_cell(r.cells[0], "Title", META_L)
    value_cell(r.cells[1],          META_MID)
    label_cell(r.cells[2], "Title", META_RL)
    value_cell(r.cells[3],          META_RV)

    # Row 9 — Signature  (taller row for wet-ink or DocuSign)
    r = t_meta.rows[9]
    label_cell(r.cells[0], "Signature", META_L)
    value_cell(r.cells[1],              META_MID)
    label_cell(r.cells[2], "Signature", META_RL)
    value_cell(r.cells[3],              META_RV)
    set_row_height(r, 500)

    # Row 10 — Date Signed | Date Approved
    r = t_meta.rows[10]
    label_cell(r.cells[0], "Date Signed",    META_L)
    value_cell(r.cells[1],                   META_MID)
    label_cell(r.cells[2], "Date Approved",  META_RL)
    value_cell(r.cells[3],                   META_RV)

    gap(doc)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 3 — Applicable To (checkboxes)
    # ──────────────────────────────────────────────────────────────────────
    # Checkbox rows:  (display_text, shade, is_subheader)
    checkbox_rows = [
        ("HealthPlan Services, Inc.  \u2610",          WHITE,      False),
        ("HealthPlan Services Insurance Agency, LLC  \u2610", WHITE, False),
        ("Policy Types",                               GRAY_SUBHDR, True),
        ("Corporate  \u2610",                          WHITE,      False),
        ("Government Affairs Review Required  \u2610", WHITE,      False),
        ("Legal Review Required  \u2610",              WHITE,      False),
        ("Line of Business (LOB)",                     GRAY_SUBHDR, True),
        ("All LOBs  \u2610",                           WHITE,      False),
        ("Specific LOB [INSERT HERE]  \u2610",         WHITE,      False),
    ]

    t_app = doc.add_table(rows=len(checkbox_rows), cols=2)
    t_app.style = "Table Grid"
    set_table_width(t_app, PAGE_W, [APP_L, APP_R])

    for i, (text, shade, is_subhdr) in enumerate(checkbox_rows):
        row = t_app.rows[i]
        cl  = row.cells[0]
        cr  = row.cells[1]

        # Left label column — "Applicable To" text only on first row
        style_cell(cl, GRAY_LABEL, APP_L)
        cl.text = ""
        pl = cl.paragraphs[0]
        set_spacing(pl, 50, 50)
        if i == 0:
            pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(pl, "Applicable To:\n(select all that apply)", bold=True, size=9.0)
        else:
            pl.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Right checkbox column
        style_cell(cr, shade, APP_R)
        cr.text = ""
        pr = cr.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(pr, 50, 50)
        run(pr, text, bold=is_subhdr, size=9.0)

    gap(doc)

    # ──────────────────────────────────────────────────────────────────────
    # HELPER — single-column section table (heading + blank content row)
    # ──────────────────────────────────────────────────────────────────────
    def section_block(heading, content_height=700):
        """
        Produces a two-row single-column table:
          Row 0: gray heading bar
          Row 1: blank white content cell with minimum height
        """
        tbl = doc.add_table(rows=2, cols=1)
        tbl.style = "Table Grid"
        set_table_width(tbl, PAGE_W, [PAGE_W])

        section_head_cell(tbl.cell(0, 0), heading, PAGE_W)

        cnt = tbl.cell(1, 0)
        style_cell(cnt, WHITE, PAGE_W)
        set_margins(cnt, top=100, bottom=100, left=120, right=120)
        cnt.text = ""
        set_spacing(cnt.paragraphs[0], 0, 0)
        set_row_height(tbl.rows[1], content_height, rule="atLeast")

        gap(doc)
        return tbl

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 4 — Purpose and Scope
    # ──────────────────────────────────────────────────────────────────────
    section_block("Purpose and Scope", content_height=1200)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 5 — Definitions
    # ──────────────────────────────────────────────────────────────────────
    section_block("Definitions", content_height=400)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 6 — Policy Statement
    # ──────────────────────────────────────────────────────────────────────
    section_block("Policy Statement", content_height=400)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 7 — Procedures
    # ──────────────────────────────────────────────────────────────────────
    section_block("Procedures", content_height=1400)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 8 — Related Policies or Standard Operating Procedures
    # ──────────────────────────────────────────────────────────────────────
    section_block("Related Policies or Standard Operating Procedures", content_height=400)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 9 — Citations / References
    # ──────────────────────────────────────────────────────────────────────
    section_block("Citations/References", content_height=400)

    # ──────────────────────────────────────────────────────────────────────
    # TABLE 10 — Revision History
    # ──────────────────────────────────────────────────────────────────────
    t_rev = doc.add_table(rows=3, cols=4)
    t_rev.style = "Table Grid"
    set_table_width(t_rev, PAGE_W, [REV_C1, REV_C2, REV_C3, REV_C4])

    # Row 0 — "Revision History" heading spanning all 4 columns
    t_rev.rows[0].cells[0].merge(t_rev.rows[0].cells[3])
    section_head_cell(t_rev.cell(0, 0), "Revision History", PAGE_W)

    # Row 1 — Column headers
    col_headers = [
        ("Date",                  REV_C1),
        ("Version Number",        REV_C2),
        ("Updated By",            REV_C3),
        ("Description of Update", REV_C4),
    ]
    for col_idx, (hdr_text, width) in enumerate(col_headers):
        c = t_rev.rows[1].cells[col_idx]
        style_cell(c, GRAY_LABEL, width)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, 55, 55)
        run(p, hdr_text, bold=True, size=9.0)

    # Row 2 — One blank data row
    for col_idx, width in enumerate([REV_C1, REV_C2, REV_C3, REV_C4]):
        c = t_rev.rows[2].cells[col_idx]
        style_cell(c, WHITE, width)
        c.text = ""
        set_spacing(c.paragraphs[0], 0, 0)
    set_row_height(t_rev.rows[2], 360)

    # ──────────────────────────────────────────────────────────────────────
    # FOOTER
    # ──────────────────────────────────────────────────────────────────────
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    set_spacing(fp, 0, 0)
    run(fp,
        "Confidential & Proprietary \u00a9 HealthPlan Services Inc., "
        "including its subsidiaries and affiliates",
        size=7.5, color="555555")

    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(fp2, 0, 0)
    run(fp2, "SEC-P[###]", size=7.5, color="555555")

    # ──────────────────────────────────────────────────────────────────────
    # SAVE
    # ──────────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"\n\u2705  Blank template saved: {output_path}\n")


# ── Entry point ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_blank_template("SEC-P_Template_Master-BLANK.docx")
